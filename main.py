from flask import Flask, request
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create the Flask app.
app = Flask(__name__)

# Set the route to 'gen-docs' and only accept POST requests
@app.route('/create-documents', methods=['POST'])
def generateDocuments():
    # Create document variables to hold processing status.
    # By default they are set to not being requested.
    # If the document is requested it will be processed successfully or unsuccessfully.
    # The string will be updated either way.
    adjudicationForm = 'document was not requested.'
    announcingSheet = 'document was not requested.'
    certificates = 'document was not requested.'
    masterJudgeRepertoire = 'document was not requested.'
    masterJudgeSchedule = 'document was not requested.'
    resultsSheet = 'document was not requested.'
    roomSchedules = 'document was not requested.'
    sessionAssignments = 'document was not requested.'
    sessionLabels = 'document was not requested.'
    teacherMaster = 'document was not requested.'

    # Read the incoming JSON data and store it in 'data'
    reqData = request.get_json()

    # Determine if the JSON object has data, if it does...process it.
    try:
        if (reqData is not None):
                # -----------------
                # Adjudication Form
                # -----------------
            if (reqData['documents']['adjudicationForm'] == True):
                adjudicationForm = createAdjudicationForm(reqData)
                # ----------------
                # Announcing Sheet
                # ----------------
            if (reqData['documents']['announcingSheet'] == True):
                announcingSheet = createAnnouncingSheet(reqData)
                # ---------------------
                # Festival Certificates
                # ---------------------
            if (reqData['documents']['certificates'] == True):
                festivalCertificates = createCertificates(reqData)
                # -----------------------
                # Master Judge Repertoire
                # -----------------------
            if (reqData['documents']['masterJudgeRepertoire'] == True):
                masterJudgeRepertoire = createMasterJudgeRepertoire(reqData)
                # ---------------------
                # Master Judge Schedule
                # ---------------------
            if (reqData['documents']['masterJudgeSchedule'] == True):
                masterJudgeSchedule = createMasterJudgeSchedule(reqData)
                # -------------
                # Results Sheet
                # -------------
            if (reqData['documents']['resultsSheet'] == True):
                resultsSheet = createResultsSheet(reqData)
                # --------------
                # Room Schedules
                # --------------
            if (reqData['documents']['roomSchedules'] == True):
                roomSchedules = createRoomSchedules(reqData)
                # -------------------
                # Session Assignments
                # -------------------
            if (reqData['documents']['sessionAssignments'] == True):
                sessionAssignments = createSessionAssignments(reqData)
                # --------------
                # Session Labels
                # --------------
            if (reqData['documents']['sessionLabels'] == True):
                sessionLabels = createSessionLabels(reqData)
                # --------------
                # Teacher Master
                # --------------
            if (reqData['documents']['teacherMaster'] == True):
                teacherMaster = createTeacherMaster(reqData)
    except Exception as err:
        return f'Something Went Wrong \n Error Details Below \n, {err}'
    else:
        return f'''\n
                   01. Adjudication Form: {adjudicationForm} \n
                   02. Announcing Sheet: {announcingSheet} \n
                   03. Festival Certificates: {certificates} \n
                   04. Master Judge Repertoire: {masterJudgeRepertoire} \n
                   05. Master Judge Schedule: {masterJudgeSchedule} \n
                   06. Results Sheet: {resultsSheet} \n
                   07. Room Schedules: {roomSchedules} \n
                   08. Session Assignments: {sessionAssignments} \n
                   09. Session Labels: {sessionLabels} \n
                   10. Teacher Master: {teacherMaster} \n
        '''

def createAdjudicationForm(data):
    try:
        #
        #
        # This document still needs to be built...
        #
        #
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createAnnouncingSheet(data):
    try:
        # # Create Document
        document = Document()

        for i in data['announcingSheet']:
            if i == 'eventName':
                eventName = data['announcingSheet'].get(i, 'Entry not found.')
            if i == 'documentTitle':
                documentTitle = data['announcingSheet'].get(i, 'Entry not found.')
            if i == 'sessions':
                for j in data['announcingSheet']['sessions']:
                    if j == 'friday':
                        # Get all session data for Friday.
                        friday = data['announcingSheet']['sessions'].get(j, 'No session data for Friday.')
                        # Process session data for each session taking place on Friday.
                        for k in range(len(friday)):
                            # Gather Student-Independent Data
                            day = str(friday[k]['day'])
                            time = str(friday[k]['time'])
                            sessionNumber = str(friday[k]['sessionNumber'])
                            nameOfRoom = str(friday[k]['nameOfRoom'])
                            location = day + ', ' + time + ', ' + sessionNumber + ', ' + nameOfRoom
                            classType = friday[k]['classType']
                            levels = friday[k]['levels']
                            judges = friday[k]['judges']
                            judge1 = judges[0].get('firstJudge', '')
                            judge2 = judges[0].get('secondJudge', '')
                            judge3 = judges[0].get('thirdJudge', '')
                            students = friday[k]['students']
                            student1 = students[0].get('firstStudent', '')
                            student2 = students[0].get('secondStudent', '')
                            student3 = students[0].get('thirdStudent', '')
                            student4 = students[0].get('fourthStudent', '')
                            student5 = students[0].get('fifthStudent', '')
                            proctorName = friday[k]['proctorName']
                            doorMonitorName = friday[k]['doorMonitorName']
                            performanceOrder = friday[k]['performanceOrder']

                            # Records Data For Table
                            records = ()

                            # Event Title
                            document.add_heading(eventName, level=0)
                            # Title Subheading
                            document.add_heading(documentTitle, level=1)
                            # Day, Time, Session Number, Room
                            document.add_heading(location, level=2)
                            # Add Performance Class Type
                            document.add_heading(classType, level=1)
                            # Add Students Levels
                            document.add_heading(f'Levels: {levels}', level=2)
                            # Add Judges
                            document.add_heading(f'Judge: {judge1}', level=3)
                            if judge2 != '':
                                document.add_heading(f'Judge: {judge2}', level=3)
                            if judge3 != '':
                                document.add_heading(f'Judge: {judge3}', level=3)
                            # Add Proctor
                            document.add_heading(f'Proctor: {proctorName}', level=3)
                            # Add Door Monitor
                            document.add_heading(f'Door Monitor: {doorMonitorName}', level=3)
                            # Performance Order
                            document.add_heading(f'{performanceOrder}:\n', level=1)

                            # Student Table
                            table = document.add_table(rows=int((len(students[0]) + 1)), cols=6)
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            table.autofit = True
                            table.style = 'TableGrid'
                            headerCells = table.rows[0].cells
                            run1 = headerCells[0].paragraphs[0].add_run('\nPerformer\n')
                            run1.bold = True
                            run1.underline = True
                            run2 = headerCells[1].paragraphs[0].add_run('\nLevel\n')
                            run2.bold = True
                            run2.underline = True
                            run3 = headerCells[2].paragraphs[0].add_run('\nSong 1\n')
                            run3.bold = True
                            run3.underline = True
                            run4 = headerCells[3].paragraphs[0].add_run('\nComposer 1\n')
                            run4.bold = True
                            run4.underline = True
                            run5 = headerCells[4].paragraphs[0].add_run('\nSong 2\n')
                            run5.bold = True
                            run5.underline = True
                            run6 = headerCells[5].paragraphs[0].add_run('\nComposer 2\n')
                            run6.bold = True
                            run6.underline = True

                            if student1[0]['studentFullName'] != '':
                                # Get Student Data
                                student1 = tuple(student1[0].values())
                                records += (student1,)
                            else:
                                student1 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student1,)
                            if student2[0]['studentFullName'] != '':
                                # Get Student Data
                                student2 = tuple(student2[0].values())
                                records += (student2,)
                            else:
                                student2 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student2,)
                            if student3[0]['studentFullName'] != '':
                                # Get Student Data
                                student3 = tuple(student3[0].values())
                                records += (student3,)
                            else:
                                student3 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student3,)
                            if student4[0]['studentFullName'] != '':
                                # Get Student Data
                                student4 = tuple(student4[0].values())
                                records += (student4,)
                            else:
                                student4 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student4,)
                            if student5[0]['studentFullName'] != '':
                                # Get Student Data
                                student5 = tuple(student5[0].values())
                                records += (student5,)
                            else:
                                student5 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student5,)

                            # Add Student Data To Table
                            i = 1
                            for perf, lvl, s1, c1, s2, c2 in records:
                                rowCells = table.rows[i].cells
                                rowCells[0].text = str(perf)
                                rowCells[1].text = str(lvl)
                                rowRun1 = rowCells[2].paragraphs[0].add_run(str(s1)).italic = True
                                rowCells[3].text = str(c1)
                                rowRun2 = rowCells[4].paragraphs[0].add_run(str(s2)).italic = True
                                rowCells[5].text = str(c2)
                                i += 1

                            # Page break after each session object is processed.
                            document.add_page_break()
                    if j == 'saturday':
                        # Get all session data for Saturday.
                        saturday = data['announcingSheet']['sessions'].get(j, 'No session data for Saturday.')
                        # Process session data for each session taking place on Saturday.
                        for k in range(len(saturday)):
                            # Gather Student-Independent Data
                            day = str(saturday[k]['day'])
                            time = str(saturday[k]['time'])
                            sessionNumber = str(saturday[k]['sessionNumber'])
                            nameOfRoom = str(saturday[k]['nameOfRoom'])
                            location = day + ', ' + time + ', ' + sessionNumber + ', ' + nameOfRoom
                            classType = saturday[k]['classType']
                            levels = saturday[k]['levels']
                            judges = saturday[k]['judges']
                            judge1 = judges[0].get('firstJudge', '')
                            judge2 = judges[0].get('secondJudge', '')
                            judge3 = judges[0].get('thirdJudge', '')
                            students = saturday[k]['students']
                            student1 = students[0].get('firstStudent', '')
                            student2 = students[0].get('secondStudent', '')
                            student3 = students[0].get('thirdStudent', '')
                            student4 = students[0].get('fourthStudent', '')
                            student5 = students[0].get('fifthStudent', '')
                            proctorName = saturday[k]['proctorName']
                            doorMonitorName = saturday[k]['doorMonitorName']
                            performanceOrder = saturday[k]['performanceOrder']

                            # Records Data For Table
                            records = ()

                            # Event Title
                            document.add_heading(eventName, level=0)
                            # Title Subheading
                            document.add_heading(documentTitle, level=1)
                            # Day, Time, Session Number, Room
                            document.add_heading(location, level=2)
                            # Add Performance Class Type
                            document.add_heading(classType, level=1)
                            # Add Students Levels
                            document.add_heading(f'Levels: {levels}', level=2)
                            # Add Judges
                            document.add_heading(f'Judge: {judge1}', level=3)
                            if judge2 != '':
                                document.add_heading(f'Judge: {judge2}', level=3)
                            if judge3 != '':
                                document.add_heading(f'Judge: {judge3}', level=3)
                            # Add Proctor
                            document.add_heading(f'Proctor: {proctorName}', level=3)
                            # Add Door Monitor
                            document.add_heading(f'Door Monitor: {doorMonitorName}', level=3)
                            # Performance Order
                            document.add_heading(f'{performanceOrder}:\n', level=1)

                            # Student Table
                            table = document.add_table(rows=int((len(students[0]) + 1)), cols=6)
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            table.autofit = True
                            table.style = 'TableGrid'
                            headerCells = table.rows[0].cells
                            run1 = headerCells[0].paragraphs[0].add_run('\nPerformer\n')
                            run1.bold = True
                            run1.underline = True
                            run2 = headerCells[1].paragraphs[0].add_run('\nLevel\n')
                            run2.bold = True
                            run2.underline = True
                            run3 = headerCells[2].paragraphs[0].add_run('\nSong 1\n')
                            run3.bold = True
                            run3.underline = True
                            run4 = headerCells[3].paragraphs[0].add_run('\nComposer 1\n')
                            run4.bold = True
                            run4.underline = True
                            run5 = headerCells[4].paragraphs[0].add_run('\nSong 2\n')
                            run5.bold = True
                            run5.underline = True
                            run6 = headerCells[5].paragraphs[0].add_run('\nComposer 2\n')
                            run6.bold = True
                            run6.underline = True

                            if student1[0]['studentFullName'] != '':
                                # Get Student Data
                                student1 = tuple(student1[0].values())
                                records += (student1,)
                            else:
                                student1 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student1,)
                            if student2[0]['studentFullName'] != '':
                                # Get Student Data
                                student2 = tuple(student2[0].values())
                                records += (student2,)
                            else:
                                student2 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student2,)
                            if student3[0]['studentFullName'] != '':
                                # Get Student Data
                                student3 = tuple(student3[0].values())
                                records += (student3,)
                            else:
                                student3 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student3,)
                            if student4[0]['studentFullName'] != '':
                                # Get Student Data
                                student4 = tuple(student4[0].values())
                                records += (student4,)
                            else:
                                student4 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student4,)
                            if student5[0]['studentFullName'] != '':
                                # Get Student Data
                                student5 = tuple(student5[0].values())
                                records += (student5,)
                            else:
                                student5 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student5,)

                            # Add Student Data To Table
                            i = 1
                            for perf, lvl, s1, c1, s2, c2 in records:
                                rowCells = table.rows[i].cells
                                rowCells[0].text = str(perf)
                                rowCells[1].text = str(lvl)
                                rowRun1 = rowCells[2].paragraphs[0].add_run(str(s1)).italic = True
                                rowCells[3].text = str(c1)
                                rowRun2 = rowCells[4].paragraphs[0].add_run(str(s2)).italic = True
                                rowCells[5].text = str(c2)
                                i += 1

                            # Page break after each session object is processed.
                            document.add_page_break()

                    if j == 'sunday':
                        # Get all session data for Sunday.
                        sunday = data['announcingSheet']['sessions'].get(j, 'No session data for Sunday.')
                        # Process session data for each session taking place on Sunday.
                        for k in range(len(sunday)):
                            # Gather Student-Independent Data
                            day = str(sunday[k]['day'])
                            time = str(sunday[k]['time'])
                            sessionNumber = str(sunday[k]['sessionNumber'])
                            nameOfRoom = str(sunday[k]['nameOfRoom'])
                            location = day + ', ' + time + ', ' + sessionNumber + ', ' + nameOfRoom
                            classType = sunday[k]['classType']
                            levels = sunday[k]['levels']
                            judges = sunday[k]['judges']
                            judge1 = judges[0].get('firstJudge', '')
                            judge2 = judges[0].get('secondJudge', '')
                            judge3 = judges[0].get('thirdJudge', '')
                            students = sunday[k]['students']
                            student1 = students[0].get('firstStudent', '')
                            student2 = students[0].get('secondStudent', '')
                            student3 = students[0].get('thirdStudent', '')
                            student4 = students[0].get('fourthStudent', '')
                            student5 = students[0].get('fifthStudent', '')
                            proctorName = sunday[k]['proctorName']
                            doorMonitorName = sunday[k]['doorMonitorName']
                            performanceOrder = sunday[k]['performanceOrder']

                            # Records Data For Table
                            records = ()

                            # Event Title
                            document.add_heading(eventName, level=0)
                            # Title Subheading
                            document.add_heading(documentTitle, level=1)
                            # Day, Time, Session Number, Room
                            document.add_heading(location, level=2)
                            # Add Performance Class Type
                            document.add_heading(classType, level=1)
                            # Add Students Levels
                            document.add_heading(f'Levels: {levels}', level=2)
                            # Add Judges
                            document.add_heading(f'Judge: {judge1}', level=3)
                            if judge2 != '':
                                document.add_heading(f'Judge: {judge2}', level=3)
                            if judge3 != '':
                                document.add_heading(f'Judge: {judge3}', level=3)
                            # Add Proctor
                            document.add_heading(f'Proctor: {proctorName}', level=3)
                            # Add Door Monitor
                            document.add_heading(f'Door Monitor: {doorMonitorName}', level=3)
                            # Performance Order
                            document.add_heading(f'{performanceOrder}:\n', level=1)

                            # Student Table
                            table = document.add_table(rows=int((len(students[0]) + 1)), cols=6)
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            table.autofit = True
                            table.style = 'TableGrid'
                            headerCells = table.rows[0].cells
                            run1 = headerCells[0].paragraphs[0].add_run('\nPerformer\n')
                            run1.bold = True
                            run1.underline = True
                            run2 = headerCells[1].paragraphs[0].add_run('\nLevel\n')
                            run2.bold = True
                            run2.underline = True
                            run3 = headerCells[2].paragraphs[0].add_run('\nSong 1\n')
                            run3.bold = True
                            run3.underline = True
                            run4 = headerCells[3].paragraphs[0].add_run('\nComposer 1\n')
                            run4.bold = True
                            run4.underline = True
                            run5 = headerCells[4].paragraphs[0].add_run('\nSong 2\n')
                            run5.bold = True
                            run5.underline = True
                            run6 = headerCells[5].paragraphs[0].add_run('\nComposer 2\n')
                            run6.bold = True
                            run6.underline = True

                            if student1[0]['studentFullName'] != '':
                                # Get Student Data
                                student1 = tuple(student1[0].values())
                                records += (student1,)
                            else:
                                student1 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student1,)
                            if student2[0]['studentFullName'] != '':
                                # Get Student Data
                                student2 = tuple(student2[0].values())
                                records += (student2,)
                            else:
                                student2 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student2,)
                            if student3[0]['studentFullName'] != '':
                                # Get Student Data
                                student3 = tuple(student3[0].values())
                                records += (student3,)
                            else:
                                student3 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student3,)
                            if student4[0]['studentFullName'] != '':
                                # Get Student Data
                                student4 = tuple(student4[0].values())
                                records += (student4,)
                            else:
                                student4 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student4,)
                            if student5[0]['studentFullName'] != '':
                                # Get Student Data
                                student5 = tuple(student5[0].values())
                                records += (student5,)
                            else:
                                student5 = ('N/A', '-', '-', '-', '-', '-')
                                records += (student5,)

                            # Add Student Data To Table
                            i = 1
                            for perf, lvl, s1, c1, s2, c2 in records:
                                rowCells = table.rows[i].cells
                                rowCells[0].text = str(perf)
                                rowCells[1].text = str(lvl)
                                rowRun1 = rowCells[2].paragraphs[0].add_run(str(s1)).italic = True
                                rowCells[3].text = str(c1)
                                rowRun2 = rowCells[4].paragraphs[0].add_run(str(s2)).italic = True
                                rowCells[5].text = str(c2)
                                i += 1

                            # Page break after each session object is processed.
                            document.add_page_break()

        # # Write/Save Document
        document.save('./documents/announcing_sheet.docx')

        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createCertificates(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createMasterJudgeRepertoire(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createMasterJudgeSchedule(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createResultsSheet(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createRoomSchedules(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createSessionAssignments(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createSessionLabels(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


def createTeacherMaster(data):
    try:
        # !
        # !
        # This document still needs to be built...
        # !
        # !
        return 'document created successfully.'
    except Exception as err:
        return f'document creation failed... \n {err}'


if __name__ == '__main_':
    app.run(debug=True, port=4321) #run app in debug mode on port 4321
